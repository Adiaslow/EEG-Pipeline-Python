import os.path as op
import re
import pyxdf
import numpy as np
import mne
from mne import create_info, io, channels, Epochs, events_from_annotations, compute_proj_raw, make_sphere_model, viz, make_forward_solution, transforms, setup_volume_source_space, preprocessing, Evoked
from mne.datasets import fetch_fsaverage
import matplotlib
import matplotlib.pyplot as plt

if __name__ == "__main__":

    # File I/O
    mainPath = 'C:/Users/admin/Desktop/Report Script/'  # Main Path
    client = 'PHJE'  # Client code
    date = '10282021'  # Date of collection
    dataf = op.join(mainPath, client + date + 'EO Full.xdf')  # Path joining
    streams, header = pyxdf.load_xdf(dataf)  # Load XDF data

    # Read Data
    ch_types = streams[0]["info"]["type"][0]
    ch_count = streams[0]["info"]["channel_count"][0]
    sfreq = float(streams[0]["info"]["nominal_srate"][0])
    ch_names = []
    ch_pos_x = []
    ch_pos_y = []
    ch_pos_z = []
    XDF_desc = str([streams[0]["info"]["desc"]])
    print(XDF_desc)
    XDF_desc_split = re.split("\'", XDF_desc)
    for i in range(13, len(XDF_desc_split) - 32, 26):
        ch_names.append(XDF_desc_split[i])
        ch_pos_x.append((XDF_desc_split[i + 8]))
        ch_pos_y.append((XDF_desc_split[i + 12]))
        ch_pos_z.append((XDF_desc_split[i + 16]))

    ch_pos = {}
    for j in range(0, len(ch_names), 1):
        temp_list = []
        temp_list.append(ch_pos_x[j])
        temp_list.append(ch_pos_y[j])
        temp_list.append(ch_pos_z[j])
        ch_pos[ch_names[j]] = temp_list

    ch_pos_x = np.array(ch_pos_x)
    ch_pos_y = np.array(ch_pos_y)
    ch_pos_z = np.array(ch_pos_z)

    # Initialize Data
    data = streams[0]["time_series"].T
    info = create_info(ch_names, sfreq, ch_types=ch_types.lower())
    raw = io.RawArray(data, info)
    raw = raw.pick_types(eeg=True).load_data()
    # montage = channels.make_dig_montage(ch_pos=ch_pos)
    montage = channels.make_standard_montage('standard_1020')
    compute_proj_raw(raw, start=0, stop=None, duration=1, n_eeg=1, n_jobs=1)
    raw.set_montage(montage, match_case=False, match_alias=False, on_missing='ignore')
    ref_channels = 'average'
    raw.set_eeg_reference(ref_channels=ref_channels, projection=True).apply_proj()  # Ave Ref
    info.set_montage(montage, match_case=False, match_alias=False, on_missing='ignore')
    # Raw EEG
    raw_filtered = raw.copy().filter(l_freq=0.5, h_freq=None)
    raw_filtered.plot(block=False, scalings=dict(eeg=25), remove_dc=True, title='Raw EEG').savefig(mainPath + 'Raw/' + 'Ave Ref Raw.png')
    raw.plot_sensors(block=False, show_names=True, sphere=(0, 0.015, 0, 0.085))

    # Preprocessing

    # Muscle Artifact
    threshold_muscle = 5
    annot_muscle, scores_muscle = preprocessing.annotate_muscle_zscore(raw, threshold=threshold_muscle, min_length_good=0.1, filter_freq=(110, 127), n_jobs=1)
    fig, ax = plt.subplots()
    ax.plot(raw.times, scores_muscle)
    ax.axhline(y=threshold_muscle, color='r')
    ax.set(xlabel='time, (s)', ylabel='zscore', title='Muscle activity')
    # fig.savefig(mainPath + 'Artifacting/' + 'Muscle Artifact.png')
    raw_filtered.set_annotations(annot_muscle)
    raw_filtered.plot(scalings=dict(eeg=25), remove_dc=True, title='Raw EEG')

    # Eye Blink Artifact
    blink_ch = [ch_names[0], ch_names[1], ch_names[10], ch_names[11]]
    blink_picks = mne.pick_channels(ch_names, blink_ch, exclude=[], ordered=False)
    threshold_eog = None
    eog_events = mne.preprocessing.find_eog_events(raw_filtered, l_freq=1, h_freq=10, filter_length='10s', ch_name=blink_ch, tstart=0, verbose=None, thresh=threshold_eog)
    onsets = eog_events[:, 0] / raw.info['sfreq'] - 0.25
    durations = [1.5] * len(eog_events)
    descriptions = ['BAD_blink', ] * len(eog_events)
    annot_eog = mne.Annotations(onsets, durations, descriptions, orig_time=raw.info['meas_date'])
    raw_filtered.set_annotations(annot_muscle + annot_eog)
    eog_epochs = mne.Epochs(raw_filtered, eog_events)
    raw_filtered.plot(events=eog_events, block=False, scalings=dict(eeg=25), remove_dc=True, highpass=0.5, lowpass=40)
    artifacted_epochs = mne.make_fixed_length_epochs(raw_filtered,reject_by_annotation=True, duration=1, preload=True)
    # artifacted_epochs.plot(block=True, scalings=dict(eeg=25))

    # ICA
    ica = preprocessing.ICA(method='fastica', n_components=None, random_state=96)
    ica.fit(artifacted_epochs)
    # https://labeling.ucsd.edu/tutorial/labels
    raw.load_data()
    ica.plot_sources(raw).savefig(mainPath + 'ICA/' + 'Raw ICA.png')
    comps = ica.plot_components(picks=None, show=False, title='ICA Components', cmap='turbo', colorbar=True, sphere=(0, 0.015, 0, 0.085))
    comps[0].savefig(mainPath + 'ICA/' + 'Components.png')
    ic = ica.plot_properties(raw, picks=range(0, ica.n_components_), show=False, dB=True, plot_std=True, psd_args={'fmax': 50}, topomap_args={'cmap': 'turbo', 'sphere': (0, 0.015, 0, 0.085)})
    cleaned_epochs = ica.apply(artifacted_epochs, include=None, exclude=None, n_pca_components=None, start=None, stop=None, verbose=None)

    for k in range(0, len(ic), 1):
        ic[k].savefig(mainPath + 'ICA/' + 'IC ' + str(k) + '.png')


    # Source Localization
    sphere = make_sphere_model(r0=(0.0, 0.0, 0.04), head_radius=0.09, info=ica.info, relative_radii=(0.9, 0.92, 0.97, 1.0), sigmas=(0.33, 1.0, 0.004, 0.33), verbose=None)
    fs_dir = fetch_fsaverage(verbose=True)
    subjects_dir = op.dirname(fs_dir)
    subject = 'fsaverage'
    trans = 'fsaverage'
    src = op.join(fs_dir, 'bem', 'fsaverage-ico-5-src.fif')
    bem = op.join(fs_dir, 'bem', 'fsaverage-5120-5120-5120-bem-sol.fif')
    viz.set_3d_backend('pyvistaqt')
    fwd = make_forward_solution(ica.info, trans, src, sphere, eeg=True, mindist=0.0, ignore_ref=False, n_jobs=1, verbose=None)
    source_space_fig = mne.viz.plot_alignment(ica.info, src=src, eeg=['original', 'projected'], trans=trans, show_axes=True, mri_fiducials=True, dig='fiducials')
    print(f'Before: {src}')
    print(f'After:  {fwd["src"]}')
    leadfield = fwd['sol']['data']
    print("Leadfield size : %d sensors x %d dipoles" % leadfield.shape)

    fwd_fixed = mne.convert_forward_solution(fwd, surf_ori=True, force_fixed=True, use_cps=True)
    leadfield = fwd_fixed['sol']['data']
    print("Leadfield size : %d sensors x %d dipoles" % leadfield.shape)

    evoked = cleaned_epochs.average()
    dip = mne.fit_dipole(evoked, bem=bem, trans=trans, cov=None, min_dist=5.0, n_jobs=1, pos=None, ori=None, rank=None, accuracy='normal', tol=5e-05, verbose=None)[0]
    dip.plot_locations(trans, subject, subjects_dir, mode='orthoview', block=True)


    '''
    from docx import Document
from docx.shared import Inches
import os, os.path as op
import re
import numpy as np
import mne
from mne import setup_volume_source_space, setup_source_space, make_forward_solution
from mne.minimum_norm import make_inverse_operator, apply_inverse_epochs
from mne_connectivity import spectral_connectivity
from mne_connectivity.viz import circular_layout, plot_connectivity_circle
import pyxdf
import matplotlib, matplotlib.pyplot as plt
matplotlib.use('Qt5Agg')
import autoreject

if __name__ == "__main__":

    # File Input
    mainPath = 'C:/Users/admin/Desktop/Report Script/'  # Main Path
    cliCod = 'PHJE'  # Client code
    datOfCol = '10282021'  # Date of collection
    dataf = os.path.join(mainPath, cliCod + datOfCol + 'EO Full.xdf')  # Path joining
    streams, header = pyxdf.load_xdf(dataf)  # Load XDF data

    # Data Extraction
    ch_types = streams[0]["info"]["type"][0]
    ch_count = streams[0]["info"]["channel_count"][0]
    sfreq = float(streams[0]["info"]["nominal_srate"][0])
    ch_names = []
    ch_pos_x = []
    ch_pos_y = []
    ch_pos_z = []
    desc = str([streams[0]["info"]["desc"]])
    print(desc)
    descS = re.split("\'", desc)
    for i in range(13, len(descS) - 32, 26):
        ch_names.append(descS[i])
        ch_pos_x.append((descS[i + 8]))
        ch_pos_y.append((descS[i + 12]))
        ch_pos_z.append((descS[i + 16]))

    ch_pos = {}
    for j in range(0, len(ch_names), 1):
        temp_list = []
        temp_list.append(ch_pos_x[j])
        temp_list.append(ch_pos_y[j])
        temp_list.append(ch_pos_z[j])
        ch_pos[ch_names[j]] = temp_list

    # Print info
    print()
    print('*****Info*****')
    print('File Path= ' + str(dataf))
    print('Data Type = ' + str(ch_types))
    print('Channel Count = ' + str(ch_count))
    print('Sample Rate = ' + str(sfreq))
    print("Channel Labels = " + str(ch_names))
    # print("Location X = " + str(ch_pos_x))
    # print("Location y = " + str(ch_pos_y))
    # print("Location Z = " + str(ch_pos_z))

    # Initializing Data
    data = streams[0]["time_series"].T
    info = mne.create_info(ch_names, sfreq, ch_types=ch_types.lower(), verbose=None)
    raw = mne.io.RawArray(data, info)
    raw = raw.pick_types(eeg=True).load_data()
    # raw.set_montage(dig_montage, on_missing='ignore')
    montage = mne.channels.make_standard_montage('standard_1020')
    raw.set_montage(montage, match_case=False, match_alias=False, on_missing='ignore', verbose=None)
    raw.set_eeg_reference(projection=True).apply_proj()  # Ave Ref

    # Raw EEG
    raw.plot(block=False, scalings=dict(eeg=25), remove_dc=True, highpass=0.5, lowpass=70, title='Raw EEG', show=False).savefig(mainPath + 'Raw/' + 'Ave Ref Raw.png')

    epochs = mne.make_fixed_length_epochs(raw, duration=1, preload=False)

    #Compute Inverse Solution
    """
    subject = 'sample'
    labels_vol = ['Left-Amygdala', 'Right-Amygdala']
    src = setup_source_space(subject, spacing='oct6', surface='white', subjects_dir=None, add_dist=True, n_jobs=1, verbose=None)
    bem = 
    make_forward_solution(raw.info, trans='fsaverage', src, bem, eeg=True, mindist=0.0, verbose=None)
    """

    # Eye Blink Repair
    blink_ch = [ch_names[0], ch_names[1], ch_names[10], ch_names[11]]
    blink_picks = mne.pick_channels(ch_names, blink_ch, exclude=[], ordered=False)
    eog_events = mne.preprocessing.find_eog_events(raw, l_freq=1, h_freq=10, filter_length='10s', ch_name=blink_ch,
                                                   tstart=0, verbose=None)
    onsets = eog_events[:, 0] / raw.info['sfreq'] - 0.25
    durations = [1.5] * len(eog_events)
    descriptions = ['blink pre ICA'] * len(eog_events)
    blink_annot_pre = mne.Annotations(onsets, durations, descriptions, orig_time=raw.info['meas_date'])
    raw.set_annotations(blink_annot_pre)
    eog_epochs = mne.Epochs(raw, eog_events)
    eog_epochs.average().plot_joint(show=False).savefig(mainPath + 'Artifacting/' + 'Blinks Raw Plot Joint.png')

    # ICA
    filt_raw = raw.copy().load_data().filter(l_freq=1., h_freq=40)
    ica = mne.preprocessing.ICA(method='picard', n_components=None, random_state=96, max_iter='auto')
    ica.fit(filt_raw)
    ica
    # https://labeling.ucsd.edu/tutorial/labels
    raw.load_data()
    ica.plot_sources(raw, show=False, block=False).savefig(mainPath + 'ICA/' + 'Raw ICA.png')
    comps = ica.plot_components(picks=None, show=False, title=None, cmap='turbo', colorbar=True, sphere=(0, 0.015, 0, 0.085))
    comps[0].savefig(mainPath + 'ICA/' + 'Components.png')
    ic = ica.plot_properties(raw, picks=range(0, ica.n_components_), dB=True, plot_std=True, show=False, verbose=None, psd_args={'fmax':70}, topomap_args={'cmap':'turbo','sphere': (0, 0.015, 0, 0.085)})
    for k in range(0, len(ic), 1):
        ic[k].savefig(mainPath + 'ICA/' + 'IC ' + str(k) + '.png')

    # EOG Removal
    ica.exclude = []
    eog_indicies, eog_scores = ica.find_bads_eog(raw, ch_name='Fp1')
    ica.exclude = eog_indicies
    ica.plot_scores(eog_scores).savefig(mainPath + 'ICA/' + 'EOG Matches.png')
    ica.plot_overlay(raw, exclude=eog_indicies, show=False).savefig(mainPath + 'ICA/' + 'EOG Component Overlay.png')

    raw_post_ICA = ica.apply(raw, exclude=eog_indicies)

    threshold_muscle = 6
    muscle_annot, scores_muscle = mne.preprocessing.annotate_muscle_zscore(raw_post_ICA, ch_type="eeg", threshold=threshold_muscle, min_length_good=0.2,filter_freq=[110, 127])
    raw_post_ICA.set_annotations(muscle_annot)
    fig, ax = plt.subplots()
    ax.plot(raw_post_ICA.times, scores_muscle)
    ax.axhline(y=threshold_muscle, color='r')
    ax.set(xlabel='time, (s)', ylabel='zscore', title='Muscle activity')

    # Eye Blink Repair Post ICA
    blink_ch = [ch_names[0], ch_names[1], ch_names[10], ch_names[11]]
    blink_picks = mne.pick_channels(ch_names, blink_ch, exclude=[], ordered=False)
    threshold_eog = 17
    eog_events_post = mne.preprocessing.find_eog_events(raw_post_ICA, l_freq=1, h_freq=10, filter_length='10s', ch_name=blink_ch, tstart=0, verbose=None, thresh=threshold_eog)
    onsets = eog_events_post[:, 0] / raw_post_ICA.info['sfreq'] - 0.25
    durations = [0.5] * len(eog_events_post)
    descriptions = ['blink post ICA'] * len(eog_events_post)
    blink_annot_post = mne.Annotations(onsets, durations, descriptions, orig_time=raw_post_ICA.info['meas_date'])
    raw_post_ICA.set_annotations(blink_annot_post)
    eog_epochs = mne.Epochs(raw_post_ICA, eog_events_post)
    eog_epochs.average().plot_joint(show=False).savefig(mainPath + 'Artifacting/' + 'Blinks Post ICA Plot Joint.png')
    raw_post_ICA.plot(events=eog_events_post, block=True, scalings=dict(eeg=25), remove_dc=True, highpass=0.5, lowpass=40)


    # Auto Rejection
    epochs_post = mne.make_fixed_length_epochs(raw_post_ICA, duration=1, preload=False)
    """
    epochs_post_orig = epochs_post.copy()
    reject = autoreject.get_rejection_threshold(epochs_post)
    print('The rejection dictionary is %s' % reject)
    epochs_post.drop_bad(reject=reject)
    bad_idx = [idx for idx, ch_log in enumerate(epochs_post.drop_log) if ch_log != ()]
    """

    # Absolute Power Maps
    raw_post_ICA.plot_psd_topomap(
        bands=[(0, 4, 'Delta'), (4, 8, 'Theta'), (8, 12, 'Alpha')], \
        ch_type='eeg', dB=True, normalize=True, cmap='turbo', show=False, verbose=None, sphere=(0, 0.015, 0, 0.085)).savefig(
        mainPath + 'Topo/' + 'Delta Theta Alpha.png')
    raw_post_ICA.plot_psd_topomap(
        bands=[(8, 10, 'Alpha 1'), (10, 12, 'Alpha 2'), (12, 35, 'Beta')], \
        ch_type='eeg', dB=True, normalize=True, cmap='turbo', show=False, verbose=None, sphere=(0, 0.015, 0, 0.085)).savefig(
        mainPath + 'Topo/' + 'Alpha 1 Alpha 2 Beta.png')
    raw_post_ICA.plot_psd_topomap(
        bands=[(12, 15, 'Beta 1'), (15, 18, 'Beta 2'), (18, 25, 'Beta 3')], \
        ch_type='eeg', dB=True, normalize=True, cmap='turbo', show=False, verbose=None, sphere=(0, 0.015, 0, 0.085)).savefig(
        mainPath + 'Topo/' + 'Beta 1 Beta 2 Beta 3.png')
    raw_post_ICA.plot_psd_topomap(
        bands=[(25, 30, 'High Beta'), (30, 45, 'Gamma 1'), (45, 80, 'Gamma 2')], \
        ch_type='eeg', dB=True, normalize=True, cmap='turbo', show=False, verbose=None, sphere=(0, 0.015, 0, 0.085)).savefig(
        mainPath + 'Topo/' + 'High Beta Gamma 1 Gamma 2.png')

    # Global PSDs
    raw_post_ICA.plot_psd(fmin=0, fmax=50, estimate='power', dB=True, average=False, spatial_colors=False, sphere=(0, 0.015, 0, 0.085), show=False, verbose=None).savefig(mainPath + '/PSD/' + 'black-PSD.png')
    raw_post_ICA.plot_psd(fmin=0, fmax=50, estimate='power', dB=True, average=True, spatial_colors=False, sphere=(0, 0.015, 0, 0.085), show=False, verbose=None).savefig(mainPath + '/PSD/' + 'average-PSD.png')
    raw_post_ICA.plot_psd(fmin=0, fmax=50, estimate='power', dB=True, average=False, spatial_colors=True, sphere=(0, 0.015, 0, 0.085), show=True, verbose=None).savefig(mainPath + '/PSD/' + 'color-PSD.png')

    # Channel PSDs
    for l in range(0, len(ch_names), 1):
        raw_post_ICA.plot_psd(fmin=0, fmax=50, estimate='power', dB=True, picks=l, show=False, verbose=None).savefig(mainPath + 'PSD/' + str(ch_names[l]) + ' PSD.png')

    """document = Document()

    document.add_heading('Comprehensive Report', 0)

    document.add_heading('Raw Channel Activity', 2)
    document.add_picture(mainPath + 'Raw/' + 'Ave Ref Raw.png', width=Inches(6))
    document.add_page_break()

    # Topo Maps
    document.add_heading('Power Topography', 1)
    document.add_heading('Average Reference', 2)
    document.add_picture(mainPath + 'Topo/' + 'Delta Theta Alpha.png', width=Inches(6))
    document.add_picture(mainPath + 'Topo/' + 'Alpha 1 Alpha 2 Beta.png', width=Inches(6))
    document.add_picture(mainPath + 'Topo/' + 'Beta 1 Beta 2 Beta 3.png', width=Inches(6))
    document.add_picture(mainPath + 'Topo/' + 'High Beta Gamma 1 Gamma 2.png', width=Inches(6))
    document.add_page_break()

    # PSDs
    document.add_heading('Power Spectral Density', 1)
    document.add_picture(mainPath + 'PSD/' + 'black-PSD.png', width=Inches(6))
    document.add_heading('Average Power Spectral Density', 2)
    document.add_picture(mainPath + 'PSD/' + 'average-PSD.png', width=Inches(6))
    document.add_heading('Spatial Color Power Spectral Density', 2)
    document.add_picture(mainPath + 'PSD/' + 'color-PSD.png', width=Inches(6))
    document.add_page_break()

    # Channel PSDs

    document.add_heading('Channel Power Spectral Density', 1)
    for m in range(0, len(ch_names), 1):
        document.add_heading(ch_names[m] + 'Power Spectral Density', 2)
        document.add_picture(mainPath + 'PSD/' + ch_names[m] + ' PSD.png', width=Inches(6))

    document.add_page_break()

    # ICA
    document.add_heading('Independent Component Analysis', 1)
    document.add_heading('Sample Component Activity', 2)
    document.add_picture(mainPath + 'ICA/' + 'Raw ICA.png', width=Inches(6))
    document.add_page_break()

    document.add_heading('Independent Components', 2)
    document.add_picture(mainPath + 'ICA/' + 'Components.png', width=Inches(6))

    for n in range(0, 19, 1):
        document.add_heading('Independent Component ' + str(n), 2)
        document.add_picture(mainPath + 'ICA/' + 'IC ' + str(n) + '.png', width=Inches(5))

    document.save(mainPath + 'Comprehensive Report.docx')
    
    """'''